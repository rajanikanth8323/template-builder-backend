# src/api/import_template.py
# -----------------------------------------------------------------
# Complete rewrite — fixes all import bugs:
#   PDF: font-size based heading detection (not istitle/endswith)
#   DOCX: direct paragraph/table iteration (no O(n2) matching)
#   Table: binding="" not "{{}}"
#   DB: includes supported_locales + tags + audit event
#   URL: handles Google Drive virus-warning confirmation page
# -----------------------------------------------------------------
import io
import json
import logging
import re
import uuid
from typing import Dict, List, Optional, Tuple

import httpx
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncEngine

router = APIRouter()
logger = logging.getLogger(__name__)


def get_engine(request: Request) -> AsyncEngine:
    engine = getattr(request.app.state, "engine", None)
    if engine is None:
        raise HTTPException(status_code=500, detail="DB engine not initialized")
    return engine


# =================================================================
# Block factory helpers
# =================================================================

def make_text_block(content: str) -> Dict:
    return {"block_id": str(uuid.uuid4()), "type": "text", "content": content.strip()}


def make_section_block(title: str) -> Dict:
    return {"block_id": str(uuid.uuid4()), "type": "section", "content": title.strip()}


def make_table_block(headers: List[str], rows: List[List[str]]) -> Dict:
    # binding="" so imported tables show empty binding cells user fills in
    columns = [{"header": h, "binding": ""} for h in headers]
    return {
        "block_id": str(uuid.uuid4()),
        "type": "table",
        "columns": columns,
        "rows": rows,
        "repeat": "",
    }


def make_image_block(src: str) -> Dict:
    return {"block_id": str(uuid.uuid4()), "type": "image", "src": src}


def _block_has_content(block: Dict) -> bool:
    if block["type"] in ("text", "section"):
        return bool((block.get("content") or "").strip())
    if block["type"] == "table":
        return bool(block.get("columns"))
    if block["type"] == "image":
        return bool((block.get("src") or "").strip())
    return True


# =================================================================
# PDF PARSER — font-size based heading detection
# =================================================================

def _collect_body_font_size(words: List[dict]) -> float:
    from collections import Counter
    sizes = [round(w.get("size", 0) or 0) for w in words if (w.get("size") or 0) > 0]
    if not sizes:
        return 0.0
    common = Counter(sizes).most_common(1)
    return float(common[0][0]) if common else 0.0


def _is_heading_by_font(word_data: dict, body_font_size: float) -> bool:
    size = word_data.get("size", 0) or 0
    if body_font_size > 0 and size >= body_font_size * 1.15:
        return True
    fontname = (word_data.get("fontname") or "").lower()
    text_val = word_data.get("text", "")
    is_bold = "bold" in fontname or "black" in fontname
    is_short = len(text_val) < 60
    has_sentence_punct = bool(re.search(r'[.!?,;]', text_val))
    if is_bold and is_short and not has_sentence_punct:
        return True
    return False


def parse_pdf_to_blocks(file_bytes: bytes) -> List[Dict]:
    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="pdfplumber not installed. Add 'pdfplumber' to requirements.txt"
        )

    blocks: List[Dict] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Pass 1: find body font size across whole document
            all_words = []
            for page in pdf.pages:
                all_words.extend(page.extract_words(extra_attrs=["size", "fontname"]) or [])
            body_font_size = _collect_body_font_size(all_words)
            logger.info(f"PDF body font size: {body_font_size}pt")

            for page_num, page in enumerate(pdf.pages):
                # ----------------------------------------------------------------
                # Build a unified list of "items" — each with a vertical position
                # so we can sort everything into true reading order at the end.
                # item = {"top": float, "type": "table"|"text"|"section", ...}
                # ----------------------------------------------------------------
                page_items = []   # collected items for this page, sorted by top
                table_bboxes = [] # bbox of every table on the page

                # --- Step 1: find all tables and record their bboxes + top pos ---
                try:
                    found_tables = page.find_tables() or []
                    for tset in found_tables:
                        if hasattr(tset, "bbox"):
                            table_bboxes.append(tset.bbox)

                    # Extract table content in the same order as find_tables()
                    for tset in found_tables:
                        try:
                            table = tset.extract()
                        except Exception:
                            continue
                        if not table:
                            continue

                        # Clean rows
                        clean_rows = []
                        for row in table:
                            cells = [str(c or "").strip() for c in row]
                            if any(cells):
                                clean_rows.append(cells)
                        if not clean_rows:
                            continue

                        first_row = clean_rows[0]
                        remaining = clean_rows[1:]
                        tbl_top = tset.bbox[1] if hasattr(tset, "bbox") else 0

                        page_items.append({
                            "top": tbl_top,
                            "kind": "table",
                            "block": make_table_block(first_row, remaining),
                        })

                except Exception as te:
                    logger.warning(f"Table extraction failed p{page_num}: {te}")

                # --- Step 2: extract words, skip those inside tables ---
                try:
                    words = page.extract_words(extra_attrs=["size", "fontname"]) or []

                    def in_table(w):
                        wx, wy = w.get("x0", 0), w.get("top", 0)
                        for bbox in table_bboxes:
                            x0, y0, x1, y1 = bbox
                            # small tolerance of 2pt
                            if (x0 - 2) <= wx <= (x1 + 2) and (y0 - 2) <= wy <= (y1 + 2):
                                return True
                        return False

                    words = [w for w in words if not in_table(w)]

                    if words:
                        # Group words into lines by vertical position (2pt bucket)
                        lines_by_top: Dict[float, List[dict]] = {}
                        for w in words:
                            top = round((w.get("top", 0) or 0) / 2) * 2
                            lines_by_top.setdefault(top, []).append(w)

                        for top in sorted(lines_by_top.keys()):
                            line_words = sorted(lines_by_top[top], key=lambda w: w.get("x0", 0))
                            line_text = " ".join(w.get("text", "") for w in line_words).strip()
                            if not line_text:
                                continue

                            # Majority vote: heading?
                            heading_votes = sum(
                                1 for w in line_words
                                if _is_heading_by_font(w, body_font_size)
                            )
                            is_heading = (
                                heading_votes > len(line_words) * 0.5
                                and len(line_text) < 120
                                and not line_text.endswith(".")
                            )
                            # ALL CAPS short line = always heading
                            if line_text.isupper() and 2 < len(line_text) < 80:
                                is_heading = True

                            page_items.append({
                                "top": top,
                                "kind": "heading" if is_heading else "text",
                                "text": line_text,
                            })

                except Exception as te:
                    logger.warning(f"Text extraction failed p{page_num}: {te}")
                    plain = page.extract_text() or ""
                    if plain.strip():
                        # Fallback: append as a single text block at position 0
                        page_items.append({"top": 0, "kind": "text", "text": plain.strip()})

                # --- Step 3: sort all page items by vertical position ---
                page_items.sort(key=lambda it: it["top"])

                # --- Step 4: emit blocks in reading order ---
                # Merge consecutive text lines into paragraphs
                current_para: List[str] = []

                def flush_para():
                    nonlocal current_para
                    if current_para:
                        content = " ".join(current_para)
                        if content.strip():
                            blocks.append(make_text_block(content))
                        current_para = []

                for item in page_items:
                    kind = item["kind"]
                    if kind == "table":
                        flush_para()
                        blocks.append(item["block"])
                    elif kind == "heading":
                        flush_para()
                        blocks.append(make_section_block(item["text"]))
                    else:  # "text"
                        current_para.append(item["text"])

                flush_para()

    except Exception as e:
        logger.error(f"PDF parse failed: {e}")
        raise HTTPException(status_code=422, detail=f"PDF parse error: {e}")

    blocks = [b for b in blocks if _block_has_content(b)]
    if not blocks:
        blocks = [make_text_block("Imported PDF content — no text could be extracted")]

    logger.info(f"PDF parsed: {len(blocks)} blocks")
    return blocks


# =================================================================
# DOCX PARSER — direct iteration, no O(n2) matching
# =================================================================

def parse_docx_to_blocks(file_bytes: bytes) -> List[Dict]:
    try:
        import docx
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="python-docx not installed. Add 'python-docx' to requirements.txt"
        )

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        blocks: List[Dict] = []

        # O(1) table lookup by element identity
        table_map = {t._element: t for t in doc.tables}
        # O(1) paragraph lookup by element identity
        para_map = {p._element: p for p in doc.paragraphs}

        for elem in doc.element.body:
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

            if tag == "p":
                para = para_map.get(elem)
                if not para:
                    continue
                text_content = para.text.strip()
                if not text_content:
                    continue

                style_name = (para.style.name if para.style else "") or ""

                # Heading by Word style
                if any(h in style_name for h in ("Heading", "Title", "Subtitle")):
                    blocks.append(make_section_block(text_content))
                    continue

                # Entirely bold short line = section
                runs_bold = [r.bold for r in para.runs if r.text.strip()]
                if (runs_bold and all(runs_bold)
                        and len(text_content) < 80
                        and not text_content.endswith(".")):
                    blocks.append(make_section_block(text_content))
                    continue

                # ALL CAPS short line = section heading
                # e.g. "LOAN OFFER LETTER", "TERMS AND CONDITIONS"
                if (text_content.isupper()
                        and len(text_content) < 80
                        and len(text_content) > 3
                        and not text_content.endswith(".")):
                    blocks.append(make_section_block(text_content))
                    continue

                blocks.append(make_text_block(text_content))

            elif tag == "tbl":
                table = table_map.get(elem)
                if not table:
                    continue

                rows_list = []
                for row in table.rows:
                    seen = set()
                    cells = []
                    for cell in row.cells:
                        cid = id(cell._tc)
                        if cid not in seen:
                            seen.add(cid)
                            cells.append(cell.text.strip())
                    rows_list.append(cells)

                if rows_list:
                    headers = rows_list[0]
                    data_rows = rows_list[1:]
                    if headers:
                        blocks.append(make_table_block(headers, data_rows))

        blocks = [b for b in blocks if _block_has_content(b)]
        return blocks if blocks else [make_text_block("Imported DOCX content")]

    except Exception as e:
        logger.error(f"DOCX parse failed: {e}")
        raise HTTPException(status_code=422, detail=f"DOCX parse error: {e}")


# =================================================================
# HTML PARSER
# =================================================================

def parse_html_to_blocks(html_content: str) -> List[Dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        clean = re.sub(r"<[^>]+>", " ", html_content)
        clean = re.sub(r"\s+", " ", clean).strip()
        paragraphs = [p.strip() for p in clean.split("\n") if p.strip()]
        return [make_text_block(p) for p in paragraphs[:30]] or [make_text_block(clean[:3000])]

    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup(["script", "style", "meta", "link", "head"]):
        tag.decompose()

    blocks: List[Dict] = []
    seen_tables = set()

    for elem in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6",
                                "p", "table", "img", "ul", "ol",
                                "blockquote", "pre"]):
        tag = elem.name

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            text_content = elem.get_text(separator=" ", strip=True)
            if text_content:
                blocks.append(make_section_block(text_content))

        elif tag == "p":
            text_content = elem.get_text(separator=" ", strip=True)
            if text_content:
                blocks.append(make_text_block(text_content))

        elif tag in ("ul", "ol"):
            items = [li.get_text(separator=" ", strip=True)
                     for li in elem.find_all("li", recursive=False)]
            if items:
                lines = []
                for i, item in enumerate(items):
                    prefix = "•" if tag == "ul" else f"{i+1}."
                    lines.append(f"{prefix} {item}")
                blocks.append(make_text_block("\n".join(lines)))

        elif tag in ("blockquote", "pre"):
            text_content = elem.get_text(strip=True)
            if text_content:
                blocks.append(make_text_block(text_content))

        elif tag == "table":
            eid = id(elem)
            if eid in seen_tables:
                continue
            seen_tables.add(eid)
            rows_data = []
            for row in elem.find_all("tr"):
                cells = [td.get_text(separator=" ", strip=True)
                         for td in row.find_all(["th", "td"])]
                if cells:
                    rows_data.append(cells)
            if rows_data:
                blocks.append(make_table_block(rows_data[0], rows_data[1:]))

        elif tag == "img":
            src = elem.get("src", "")
            if src and not src.startswith("data:"):
                blocks.append(make_image_block(src))

    blocks = [b for b in blocks if _block_has_content(b)]
    return blocks if blocks else [make_text_block("Imported HTML content")]


# =================================================================
# URL FETCHER — handles Google Drive, Dropbox, OneDrive, direct links
# =================================================================

def _extract_google_drive_file_id(url: str) -> Optional[str]:
    """Extract file ID from any Google Drive URL format."""
    # /file/d/{id}/view  or  /file/d/{id}/edit
    m = re.search(r"/file/d/([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    # ?id={id}  or  &id={id}
    m = re.search(r"[?&]id=([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    # /open?id={id}
    m = re.search(r"/open\?id=([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    return None


def _normalize_url(url: str) -> str:
    """
    Convert share URLs from Google Drive / Dropbox / OneDrive
    into direct download URLs.
    """
    # ── Google Drive ──────────────────────────────────────────────
    if "drive.google.com" in url or "docs.google.com" in url:
        file_id = _extract_google_drive_file_id(url)
        if file_id:
            # drive.usercontent.google.com is the current working download endpoint
            # (drive.google.com/uc was deprecated and now returns 403)
            return f"https://drive.usercontent.google.com/download?id={file_id}&export=download&confirm=t"

    # ── Dropbox ───────────────────────────────────────────────────
    if "dropbox.com" in url:
        # Change dl=0 → dl=1  or add dl=1
        url = re.sub(r"[?&]dl=0", "", url)
        sep = "&" if "?" in url else "?"
        return f"{url}{sep}dl=1"

    # ── OneDrive ──────────────────────────────────────────────────
    if "1drv.ms" in url or "onedrive.live.com" in url:
        # Convert share link to direct download
        # 1drv.ms links need to be accessed directly — they redirect
        return url

    return url


def _is_google_login_page(html: str) -> bool:
    """Detect if Google redirected us to a login page instead of the file."""
    indicators = [
        "accounts.google.com",
        "Sign in - Google Accounts",
        "google.com/accounts",
        "ServiceLogin",
        "signin/v2",
    ]
    return any(ind in html for ind in indicators)


def _extract_drive_confirm_url(html: str, file_id: str) -> Optional[str]:
    """
    Parse the Google Drive virus-warning/large-file confirmation page
    and extract the confirmed download URL.
    """
    # Look for confirm= token anywhere in page
    m = re.search(r'[?&]confirm=([a-zA-Z0-9_-]+)', html)
    if m:
        token = m.group(1)
        if token != "t":  # "t" is the generic token we already tried
            return f"https://drive.usercontent.google.com/download?id={file_id}&export=download&confirm={token}"

    # Look for download form action URL
    m = re.search(r'action="(https://drive\.usercontent\.google\.com[^"]+)"', html)
    if m:
        return m.group(1).replace("&amp;", "&")

    # Fallback: look for any usercontent download link
    m = re.search(r'href="(https://drive\.usercontent\.google\.com/download[^"]+)"', html)
    if m:
        return m.group(1).replace("&amp;", "&")

    return None


async def fetch_url_content(url: str) -> Tuple[bytes, str]:
    """
    Fetch file content from any public URL.
    Handles: direct links, Google Drive, Dropbox, OneDrive.
    """
    original_url = url
    url = _normalize_url(url)
    logger.info(f"Fetching URL: {url}")

    # Browser-like headers to avoid 403s from servers that block bots
    hdrs = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Accept-Encoding": "gzip, deflate, br",
    }

    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=60,
            verify=False,
        ) as client:

            resp = await client.get(url, headers=hdrs)
            ct = resp.headers.get("content-type", "").lower()
            body_preview = resp.content[:200].decode("utf-8", errors="ignore")

            # ── Detect proxy/network blockage ──────────────────────────
            # When Docker has no internet, egress proxy returns:
            # HTTP 403  Content-Type: text/plain  Body: "Host not in allowlist"
            if "host not in allowlist" in body_preview.lower():
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "The server cannot reach this URL. "
                        "Please download the file to your computer and upload it directly "
                        "using the 'Upload File' tab instead."
                    )
                )

            if resp.status_code == 404:
                raise HTTPException(
                    status_code=400,
                    detail="File not found (HTTP 404). Check the URL is correct."
                )
            if resp.status_code == 403:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Access denied (HTTP 403). "
                        "For Google Drive: File → Share → 'Anyone with the link' → Viewer. "
                        "Or download the file and use 'Upload File' instead."
                    )
                )
            if resp.status_code != 200:
                raise HTTPException(
                    status_code=400,
                    detail=f"Failed to fetch URL (HTTP {resp.status_code})."
                )

            # ── Google Drive special handling ──────────────────────────
            is_drive_url = "drive.google.com" in url or "drive.usercontent.google.com" in url
            if is_drive_url:
                file_id = _extract_google_drive_file_id(original_url)

                # Got a login/permission page
                if "text/html" in ct and file_id and _is_google_login_page(resp.text):
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            "Google Drive requires sign-in for this file. "
                            "Go to Drive → Share → 'Anyone with the link' → Viewer → Done. "
                            "Or download the file and use 'Upload File' instead."
                        )
                    )

                # Got a virus-warning confirmation page — extract real download link
                if "text/html" in ct and file_id:
                    confirm_url = _extract_drive_confirm_url(resp.text, file_id)
                    if confirm_url:
                        logger.info(f"Drive confirmation page, retrying: {confirm_url}")
                        resp = await client.get(confirm_url, headers=hdrs)
                        ct = resp.headers.get("content-type", "").lower()
                        if "text/html" in ct and _is_google_login_page(resp.text):
                            raise HTTPException(
                                status_code=400,
                                detail="Google Drive requires login. Set sharing to 'Anyone with the link'."
                            )

                # Still HTML — can't download
                if "text/html" in ct and file_id:
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            "Could not download from Google Drive. "
                            "Please download the file and use 'Upload File' tab instead."
                        )
                    )

            # ── Detect HTML when we expected a file ────────────────────
            # If URL doesn't end in a file extension and we got HTML back,
            # that's probably a webpage (valid for HTML import)
            # But if it looks like an error page, warn clearly
            if "text/html" in ct:
                text_lower = resp.text[:500].lower()
                error_phrases = [
                    "access denied", "permission denied", "not found",
                    "error", "forbidden", "unauthorized", "sign in",
                ]
                url_has_ext = any(url.lower().endswith(e) for e in [".pdf", ".docx", ".txt"])
                if url_has_ext and any(p in text_lower for p in error_phrases):
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            f"The URL returned an error page instead of a file. "
                            f"The file may not be publicly accessible. "
                            f"Please download the file and use 'Upload File' tab instead."
                        )
                    )

            return resp.content, ct

    except HTTPException:
        raise  # re-raise our own errors
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="URL fetch timed out (60s). Try a smaller file.")
    except httpx.RequestError as e:
        raise HTTPException(status_code=400, detail=f"Could not reach URL: {e}")


# =================================================================
# FILE TYPE DISPATCHER
# =================================================================

def _parse_file(file_bytes: bytes, filename: str, content_type: str) -> List[Dict]:
    fn = filename.lower()
    ct = content_type.lower()

    if fn.endswith(".docx") or "wordprocessingml" in ct or "officedocument.wordprocessing" in ct:
        logger.info("Parsing as DOCX")
        return parse_docx_to_blocks(file_bytes)

    if fn.endswith(".pdf") or "application/pdf" in ct:
        logger.info("Parsing as PDF")
        return parse_pdf_to_blocks(file_bytes)

    if fn.endswith((".html", ".htm")) or "text/html" in ct:
        logger.info("Parsing as HTML")
        return parse_html_to_blocks(file_bytes.decode("utf-8", errors="ignore"))

    # Unknown — detect
    try:
        text_content = file_bytes.decode("utf-8", errors="ignore")
        if "<html" in text_content.lower() or "<!doctype" in text_content.lower():
            logger.info("Detected HTML in unknown file")
            return parse_html_to_blocks(text_content)
        paragraphs = re.split(r"\n\s*\n", text_content)
        blocks = [make_text_block(p.strip()) for p in paragraphs if p.strip()]
        return blocks[:40] or [make_text_block(text_content[:3000])]
    except Exception:
        raise HTTPException(
            status_code=422,
            detail="Unsupported file format. Supported: PDF, DOCX, HTML, plain text."
        )


# =================================================================
# ENDPOINTS
# =================================================================

@router.post("/templates/import/file", response_model=dict)
async def import_template_file(
    request: Request,
    file: UploadFile = File(...),
    name: str = Form(...),
    industry: str = Form(default=""),
    output_target: str = Form(default="html"),
):
    """Import a template from an uploaded PDF, DOCX, or HTML file."""
    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large. Maximum 20MB.")

    engine = get_engine(request)
    user = request.headers.get("x-user-id", "dev_user")
    filename = file.filename or ""
    content_type = file.content_type or ""

    logger.info(f"Import file: {filename} ({len(file_bytes)}B) ct={content_type}")
    blocks = _parse_file(file_bytes, filename, content_type)

    return await _create_template(engine, name, industry, output_target, blocks, user,
                                   f"Imported from {filename}")


@router.post("/templates/import/url", response_model=dict)
async def import_template_url(
    request: Request,
    url: str = Form(...),
    name: str = Form(...),
    industry: str = Form(default=""),
    output_target: str = Form(default="html"),
):
    """Import a template from a public URL."""
    engine = get_engine(request)
    user = request.headers.get("x-user-id", "dev_user")

    logger.info(f"Import URL: {url}")
    file_bytes, content_type = await fetch_url_content(url)
    url_path = url.split("?")[0].lower()
    blocks = _parse_file(file_bytes, url_path, content_type)

    return await _create_template(engine, name, industry, output_target, blocks, user,
                                   f"Imported from URL: {url}")


# =================================================================
# DB HELPER
# =================================================================

async def _create_template(
    engine: AsyncEngine,
    name: str,
    industry: str,
    output_target: str,
    blocks: List[Dict],
    user: str,
    import_note: str,
) -> Dict:
    template_id = str(uuid.uuid4())
    layout_json = {"blocks": blocks}

    # Step 1: ensure uuid-ossp extension exists (run separately, outside transaction)
    try:
        async with engine.begin() as conn:
            await conn.execute(text('CREATE EXTENSION IF NOT EXISTS "uuid-ossp"'))
    except Exception:
        pass  # already exists or no permission — fine either way

    # Step 2: insert the template — explicit UUID cast for asyncpg compatibility
    async with engine.begin() as conn:
        await conn.execute(text("""
            INSERT INTO template_builder.templates (
                template_id, name, description, status, output_target,
                layout_json, default_locale, supported_locales, tags,
                industry, created_by, created_at, updated_at
            ) VALUES (
                CAST(:tid AS uuid), :name, :desc, 'draft', :ot,
                CAST(:layout AS jsonb), 'en', ARRAY['en'], ARRAY[]::text[],
                :industry, :user, NOW(), NOW()
            )
        """), {
            "tid": template_id,
            "name": name,
            "desc": import_note,
            "ot": output_target,
            "layout": json.dumps(layout_json),
            "industry": industry or None,
            "user": user,
        })
    logger.info(f"Template row inserted OK: {template_id}")

    # Step 3: audit event — completely separate transaction, non-fatal
    try:
        async with engine.begin() as conn:
            await conn.execute(text("""
                INSERT INTO template_builder.audit_events (
                    event_id, entity_type, entity_id, action, actor,
                    summary, details_json, created_at
                ) VALUES (
                    uuid_generate_v4(), 'template', CAST(:tid AS uuid), 'import', :user,
                    :summary, CAST(:details AS jsonb), NOW()
                )
            """), {
                "tid": template_id,
                "user": user,
                "summary": f"Template imported: {name} ({len(blocks)} blocks)",
                "details": json.dumps({
                    "import_source": import_note,
                    "block_count": len(blocks),
                    "output_target": output_target,
                    "industry": industry or None,
                }),
            })
    except Exception as ae:
        logger.warning(f"Audit insert failed (non-fatal): {ae}")

    logger.info(f"Template created: {template_id} — {name} — {len(blocks)} blocks")
    return {
        "template_id": template_id,
        "name": name,
        "status": "draft",
        "block_count": len(blocks),
        "message": f"Template imported successfully with {len(blocks)} blocks",
    }