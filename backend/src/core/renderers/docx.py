# # src/core/renderers/docx.py
# # DOCX renderer using python-docx
# # Supports align, fontSize, and table repeat from context.datasets

# import json
# import re
# from typing import Any, Dict, List

# from docx import Document
# from docx.shared import Pt, RGBColor, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement


# class DocxRenderer:

#     def render(self, layout_json: Any, context: Dict[str, Any]) -> bytes:
#         if isinstance(layout_json, str):
#             layout_json = json.loads(layout_json)

#         doc = Document()
#         for section in doc.sections:
#             section.top_margin    = Inches(1)
#             section.bottom_margin = Inches(1)
#             section.left_margin   = Inches(1.2)
#             section.right_margin  = Inches(1.2)

#         style = doc.styles['Normal']
#         style.font.name = 'Calibri'
#         style.font.size = Pt(11)

#         blocks = layout_json.get("blocks", [])
#         if not blocks:
#             doc.add_paragraph("No content in this template.")
#             return self._to_bytes(doc)

#         for block in blocks:
#             btype = block.get("type", "")
#             if btype == "section":
#                 self._render_section(doc, block, context)
#             elif btype == "text":
#                 self._render_text(doc, block, context)
#             elif btype == "table":
#                 self._render_table(doc, block, context)
#             elif btype == "image":
#                 self._render_image(doc, block, context)

#         return self._to_bytes(doc)

#     def _render_section(self, doc, block, context):
#         content = self._replace_tokens(block.get("content", "Section"), context)
#         para    = doc.add_paragraph()
#         para.paragraph_format.space_before = Pt(14)
#         para.paragraph_format.space_after  = Pt(4)
#         run      = para.add_run(content.upper())
#         run.bold = True
#         run.font.size      = Pt(11)
#         run.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
#         pPr    = para._p.get_or_add_pPr()
#         pBdr   = OxmlElement('w:pBdr')
#         bottom = OxmlElement('w:bottom')
#         bottom.set(qn('w:val'),   'single')
#         bottom.set(qn('w:sz'),    '4')
#         bottom.set(qn('w:space'), '1')
#         bottom.set(qn('w:color'), 'A78BFA')
#         pBdr.append(bottom)
#         pPr.append(pBdr)

#     def _render_text(self, doc, block, context):
#         content = self._replace_tokens(block.get("content", ""), context)
#         if not content.strip():
#             return
#         align_map  = {
#             "left":   WD_ALIGN_PARAGRAPH.LEFT,
#             "center": WD_ALIGN_PARAGRAPH.CENTER,
#             "right":  WD_ALIGN_PARAGRAPH.RIGHT,
#         }
#         align_val = block.get("align", "left")
#         font_size = block.get("fontSize", 14)
#         for line in content.split("\n"):
#             if not line.strip():
#                 doc.add_paragraph()
#                 continue
#             para = doc.add_paragraph()
#             para.paragraph_format.space_after  = Pt(6)
#             para.paragraph_format.line_spacing = Pt(font_size * 1.4)
#             para.paragraph_format.alignment    = align_map.get(align_val, WD_ALIGN_PARAGRAPH.LEFT)
#             run           = para.add_run(line)
#             run.font.size = Pt(font_size)

#     def _render_table(self, doc, block, context):
#         cols     = block.get("columns", [])
#         block_id = block.get("block_id", "")

#         if not cols:
#             return

#         # Check datasets for repeat rows
#         datasets  = context.get("datasets", {})
#         data_rows = datasets.get(block_id, [])

#         if data_rows:
#             # Real dataset rows from DB
#             all_rows = data_rows
#         else:
#             # Fallback: binding row + manual rows
#             binding_row = [
#                 self._replace_tokens(col.get("binding", ""), context)
#                 for col in cols
#             ]
#             manual_rows = block.get("rows", [])
#             manual_data = [
#                 [
#                     self._replace_tokens(
#                         row[ci] if ci < len(row) else col.get("binding", ""),
#                         context
#                     )
#                     for ci, col in enumerate(cols)
#                 ]
#                 for row in manual_rows
#             ]
#             all_rows = [binding_row] + manual_data

#         doc.add_paragraph()
#         total_rows = 1 + len(all_rows)
#         table      = doc.add_table(rows=total_rows, cols=len(cols))
#         table.style = 'Table Grid'

#         # Header row
#         header_row = table.rows[0]
#         for ci, col in enumerate(cols):
#             cell      = header_row.cells[ci]
#             cell.text = col.get("header", "")
#             para      = cell.paragraphs[0]
#             run       = para.runs[0] if para.runs else para.add_run(cell.text)
#             run.bold           = True
#             run.font.size      = Pt(10)
#             run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
#             tc_pr = cell._tc.get_or_add_tcPr()
#             shd   = OxmlElement('w:shd')
#             shd.set(qn('w:val'),   'clear')
#             shd.set(qn('w:color'), 'auto')
#             shd.set(qn('w:fill'),  '4F46E5')
#             tc_pr.append(shd)

#         # Data rows
#         for ri, row_data in enumerate(all_rows):
#             table_row = table.rows[ri + 1]
#             bg        = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
#             for ci, col in enumerate(cols):
#                 cell_val  = str(row_data[ci]) if ci < len(row_data) else ""
#                 cell      = table_row.cells[ci]
#                 cell.text = cell_val
#                 para      = cell.paragraphs[0]
#                 run       = para.runs[0] if para.runs else para.add_run(cell_val)
#                 run.font.size = Pt(10)
#                 tc_pr = cell._tc.get_or_add_tcPr()
#                 shd   = OxmlElement('w:shd')
#                 shd.set(qn('w:val'),   'clear')
#                 shd.set(qn('w:color'), 'auto')
#                 shd.set(qn('w:fill'),  bg)
#                 tc_pr.append(shd)

#         doc.add_paragraph()

#     def _render_image(self, doc, block, context):
#         src = self._replace_tokens(block.get("src", ""), context)
#         if src:
#             para           = doc.add_paragraph()
#             run            = para.add_run(f"[Image: {src}]")
#             run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
#             run.font.size  = Pt(10)
#             run.italic     = True

#     def _replace_tokens(self, content: str, context: Dict[str, Any]) -> str:
#         if not content:
#             return ""
#         values = context.get("values", {})
#         for key, val in values.items():
#             content = content.replace("{{" + key + "}}", str(val) if val is not None else "")
#         content = re.sub(r'\{\{[^}]+\}\}', '', content)
#         return content.strip()

#     def _to_bytes(self, doc: Document) -> bytes:
#         import io
#         buf = io.BytesIO()
#         doc.save(buf)
#         return buf.getvalue()
# src/core/renderers/docx.py
# DOCX renderer using python-docx
# Supports align, fontSize, and table repeat from context.datasets

import json
import re
import base64
import urllib.request
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocxRenderer:

    def render(self, layout_json: Any, context: Dict[str, Any]) -> bytes:
        if isinstance(layout_json, str):
            layout_json = json.loads(layout_json)

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        blocks = layout_json.get("blocks", [])
        if not blocks:
            doc.add_paragraph("No content in this template.")
            return self._to_bytes(doc)

        for block in blocks:
            btype = block.get("type", "")
            if btype == "section":
                self._render_section(doc, block, context)
            elif btype == "text":
                self._render_text(doc, block, context)
            elif btype == "table":
                self._render_table(doc, block, context)
            elif btype == "image":
                self._render_image(doc, block, context)

        return self._to_bytes(doc)

    def _render_section(self, doc, block, context):
        content = self._replace_tokens(block.get("content", "Section"), context)
        para    = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(4)
        run      = para.add_run(content.upper())
        run.bold = True
        run.font.size      = Pt(11)
        run.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
        pPr    = para._p.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'A78BFA')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_text(self, doc, block, context):
        content = self._replace_tokens(block.get("content", ""), context)
        if not content.strip():
            return
        align_map  = {
            "left":   WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        }
        align_val = block.get("align", "left")
        font_size = block.get("fontSize", 14)
        for line in content.split("\n"):
            if not line.strip():
                doc.add_paragraph()
                continue
            para = doc.add_paragraph()
            para.paragraph_format.space_after  = Pt(6)
            para.paragraph_format.line_spacing = Pt(font_size * 1.4)
            para.paragraph_format.alignment    = align_map.get(align_val, WD_ALIGN_PARAGRAPH.LEFT)
            run           = para.add_run(line)
            run.font.size = Pt(font_size)

    def _render_table(self, doc, block, context):
        cols     = block.get("columns", [])
        block_id = block.get("block_id", "")

        if not cols:
            return

        datasets  = context.get("datasets", {})
        data_rows = datasets.get(block_id, [])

        if data_rows:
            all_rows = data_rows
        else:
            binding_row = [
                self._replace_tokens(col.get("binding", ""), context)
                for col in cols
            ]
            manual_rows = block.get("rows", [])
            manual_data = [
                [
                    self._replace_tokens(
                        row[ci] if ci < len(row) else col.get("binding", ""),
                        context
                    )
                    for ci, col in enumerate(cols)
                ]
                for row in manual_rows
            ]
            all_rows = [binding_row] + manual_data

        doc.add_paragraph()
        total_rows = 1 + len(all_rows)
        table      = doc.add_table(rows=total_rows, cols=len(cols))
        table.style = 'Table Grid'

        # Header row
        header_row = table.rows[0]
        for ci, col in enumerate(cols):
            cell      = header_row.cells[ci]
            cell.text = col.get("header", "")
            para      = cell.paragraphs[0]
            run       = para.runs[0] if para.runs else para.add_run(cell.text)
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd   = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  '4F46E5')
            tc_pr.append(shd)

        # Data rows
        for ri, row_data in enumerate(all_rows):
            table_row = table.rows[ri + 1]
            bg        = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
            for ci, col in enumerate(cols):
                cell_val  = str(row_data[ci]) if ci < len(row_data) else ""
                cell      = table_row.cells[ci]
                cell.text = cell_val
                para      = cell.paragraphs[0]
                run       = para.runs[0] if para.runs else para.add_run(cell_val)
                run.font.size = Pt(10)
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  bg)
                tc_pr.append(shd)

        doc.add_paragraph()

    def _render_image(self, doc, block, context):
        src = self._replace_tokens(block.get("src", ""), context)
        if not src:
            return

        try:
            img_bytes = self._load_image_bytes(src)
            if img_bytes:
                img_buf = BytesIO(img_bytes)
                doc.add_picture(img_buf, width=Inches(5.5))
                doc.add_paragraph()
            else:
                para = doc.add_paragraph()
                run  = para.add_run(f"[Image not found: {src}]")
                run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                run.font.size  = Pt(10)
                run.italic     = True
        except Exception as exc:
            para = doc.add_paragraph()
            run  = para.add_run(f"[Image error: {exc}]")
            run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
            run.font.size  = Pt(10)
            run.italic     = True

    def _load_image_bytes(self, src: str):
        """Return raw image bytes from a base64 data-URL or http/https URL."""
        if src.startswith("data:"):
            # e.g. data:image/png;base64,AAAA...
            _, encoded = src.split(",", 1)
            return base64.b64decode(encoded)
        elif src.startswith("http://") or src.startswith("https://"):
            req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                return resp.read()
        return None

    def _replace_tokens(self, content: str, context: Dict[str, Any]) -> str:
        if not content:
            return ""
        values = context.get("values", {})
        for key, val in values.items():
            content = content.replace("{{" + key + "}}", str(val) if val is not None else "")
        content = re.sub(r'\{\{[^}]+\}\}', '', content)
        return content.strip()

    def _to_bytes(self, doc: Document) -> bytes:
        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()